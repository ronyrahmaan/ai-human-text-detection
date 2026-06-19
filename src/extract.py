"""Pull plain text out of uploaded PDF and Word documents."""
from __future__ import annotations

import io


class ExtractionError(Exception):
    """Raised when a document yields no usable text."""


def extract_text(file_name: str, data: bytes) -> str:
    """Return the text of a .pdf, .docx, or .txt file given its raw bytes."""
    name = file_name.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")
    raise ExtractionError(f"Unsupported file type: {file_name}")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001 - surfaced to the user
            raise ExtractionError("This PDF is encrypted and cannot be read.") from exc

    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    if not text.strip():
        raise ExtractionError(
            "No text found. The PDF is likely scanned (images only); "
            "OCR would be needed to read it.")
    return text


def _from_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    text = "\n".join(p for p in parts if p)
    if not text.strip():
        raise ExtractionError("No text found in the document.")
    return text
